import streamlit as st
import pandas as pd
from openai import OpenAI
from io import BytesIO
from docx import Document

# ==========================================
# 1. 页面配置与 CSS 样式 (纯净版)
# ==========================================
st.set_page_config(page_title="Bosch VDA6.3 审核助手", page_icon="🛡️", layout="wide", initial_sidebar_state="expanded")

# 将长串 CSS 单独拿出来，避免引号冲突
custom_css = """
<style>
    .stApp::before {
        content: ""; position: fixed; top: 0; left: 0; right: 0; height: 8px; z-index: 999999;
        background: linear-gradient(90deg, #A8000B 0%, #E3000F 11%, #FFB600 22%, #FFF000 33%, #8CBE00 44%, #009933 55%, #00A6EB 66%, #006FB4 77%, #003B6A 88%, #642C90 100%);
    }
    .stApp {
        background: radial-gradient(circle at 10% 20%, #f1f5f9, #e2e8f0);
        color: #1e293b; font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }
    [data-testid="stSidebar"] {
        background-color: rgba(248, 250, 252, 0.7) !important;
        backdrop-filter: blur(20px); -webkit-backdrop-filter: blur(20px);
        border-right: 1px solid rgba(226, 232, 240, 0.8); padding-top: 1rem; 
    }
    [data-testid="stSidebarNav"] {display: none;}
    
    div.nav-btn > div > div > button {
        background-color: transparent; color: #475569; border: none; border-radius: 8px; 
        padding: 0.6rem 1rem; margin-bottom: 0.5rem; width: 100%; text-align: left; 
        font-weight: 500; font-size: 0.95rem; transition: all 0.2s ease; box-shadow: none; display: flex; justify-content: flex-start;
    }
    div.nav-btn > div > div > button:hover { background-color: rgba(15, 23, 42, 0.05); color: #0f172a; transform: translateY(0); box-shadow: none; }
    div.nav-btn.active > div > div > button { background-color: rgba(15, 23, 42, 0.08); color: #003B6A; font-weight: 700; }
    
    [data-testid="stDataFrame"], .stFileUploader, div.stInfo, div.stWarning, div.stSuccess, div[data-testid="stForm"] {
        background-color: rgba(255, 255, 255, 0.55) !important;
        backdrop-filter: blur(12px) !important; -webkit-backdrop-filter: blur(12px) !important;
        border-radius: 12px; border: 1px solid rgba(255, 255, 255, 0.8) !important;
        box-shadow: 0 8px 32px 0 rgba(31, 38, 135, 0.05); transition: all 0.3s ease;
    }
    
    .stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div>div {
        background-color: rgba(255, 255, 255, 0.5) !important; color: #0f172a;
        border: 1px solid rgba(203, 213, 225, 0.6); border-radius: 8px; transition: all 0.2s;
    }
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus, .stSelectbox>div>div>div:focus {
        background-color: rgba(255, 255, 255, 0.9) !important; border-color: #0070c0;
        box-shadow: 0 0 0 3px rgba(0, 112, 192, 0.15) !important;
    }
    
    h1, h2, h3 { color: #003B6A !important; font-weight: 700 !important; letter-spacing: -0.02em; }
    
    div.main-btn > div > div > button, .stDownloadButton>button {
        background: rgba(0, 112, 192, 0.9); backdrop-filter: blur(4px); color: #ffffff;
        border: 1px solid rgba(255, 255, 255, 0.2); border-radius: 8px; padding: 0.6rem 1.2rem;
        font-weight: 600; box-shadow: 0 4px 14px 0 rgba(0, 112, 192, 0.2);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1); width: 100%;
    }
    div.main-btn > div > div > button:hover, .stDownloadButton>button:hover {
        transform: translateY(-2px); box-shadow: 0 6px 20px rgba(0, 112, 192, 0.4); background: rgba(0, 112, 192, 1); color: white;
    }
    #MainMenu {visibility: hidden;} footer {visibility: hidden;}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# ==========================================
# 2. 会话状态与核心配置
# ==========================================
DEEPSEEK_API_KEY = "sk-a43acaccdd004f98a6c1693c47c01b50"
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

if 'vda_data' not in st.session_state: st.session_state['vda_data'] = None 
if 'clause_list' not in st.session_state: st.session_state['clause_list'] = []
if 'current_page' not in st.session_state: st.session_state['current_page'] = "page1"

if 'ai_finding_result' not in st.session_state: st.session_state['ai_finding_result'] = ""
if 'ai_score_result' not in st.session_state: st.session_state['ai_score_result'] = ""
if 'five_why_result' not in st.session_state: st.session_state['five_why_result'] = None

# ==========================================
# 3. 辅助函数定义
# ==========================================
def nav_button(label, page_key, icon=""):
    css_class = "nav-btn active" if st.session_state['current_page'] == page_key else "nav-btn"
    st.sidebar.markdown(f'<div class="{css_class}">', unsafe_allow_html=True)
    if st.sidebar.button(f"{icon} {label}", key=f"btn_{page_key}"):
        st.session_state['current_page'] = page_key
        st.rerun()
    st.sidebar.markdown('</div>', unsafe_allow_html=True)

def create_word_report(finding, five_why, person, deadline):
    doc = Document()
    doc.add_heading('VDA 6.3 问题原因分析报告 (5-Why)', 0)
    doc.add_heading('一、 审核发现 (Finding)', level=1)
    doc.add_paragraph(finding)
    doc.add_heading('二、 5-Why 根本原因推导', level=1)
    doc.add_paragraph(five_why)
    doc.add_heading('三、 改善计划', level=1)
    doc.add_paragraph(f"责任人: {person}")
    doc.add_paragraph(f"完成期限: {deadline}")
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ==========================================
# 4. 侧边栏导航渲染
# ==========================================
st.sidebar.markdown("<h2 style='margin-bottom: 0px;'>🛡️ Bosch</h2>", unsafe_allow_html=True)
st.sidebar.markdown("<p style='color: #64748b; font-size: 0.85rem; font-weight: 500; margin-top: -5px; margin-bottom: 30px;'>审核助手 Co-pilot</p>", unsafe_allow_html=True)
nav_button("导入审核标准", "page1", "⚙️")
nav_button("录入与 AI 分析", "page2", "🧠")
nav_button("问题原因分析", "page3", "📑") 

# ==========================================
# 5. 核心页面逻辑
# ==========================================
menu = st.session_state['current_page']

if menu == "page1":
    st.title("⚙️ 导入 VDA6.3 审核标准")
    if st.session_state['vda_data'] is not None:
        st.success("✅ **审核标准已成功导入！**")
        st.dataframe(st.session_state['vda_data'], height=350, use_container_width=True)
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🔄 清空并上传新文件"):
            st.session_state['vda_data'] = None
            st.session_state['clause_list'] = []
            st.rerun() 
    else:
        st.markdown("<p style='color: #475569;'>请上传最新的 VDA6.3 审核标准文件 (.xlsx 或 .csv)。</p>", unsafe_allow_html=True)
        uploaded_file = st.file_uploader("📥 点击或拖拽文件上传", type=["xlsx", "xls", "csv"])
        if uploaded_file is not None:
            try:
                if uploaded_file.name.endswith('.csv'): df = pd.read_csv(uploaded_file)
                else: df = pd.read_excel(uploaded_file)
                st.session_state['vda_data'] = df
                if "reference" in df.columns: st.session_state['clause_list'] = df["reference"].dropna().astype(str).unique().tolist()
                st.rerun()
            except Exception as e:
                st.error(f"❌ 读取文件失败：{e}")

elif menu == "page2":
    st.title("🧠 现场问题 AI 分析")
    if not st.session_state['clause_list']:
        st.warning("⚠️ **缺少审核标准数据**。请先前往左侧 [导入审核标准] 页面上传。")
    else:
        with st.container():
            col1, col2 = st.columns([1, 1.1], gap="large") 
            with col1:
                st.markdown("#### 📝 第一步：审核员记录")
                clause = st.selectbox("🎯 选择审核条款编号 (reference)", st.session_state['clause_list'])
                raw_finding = st.text_area("✍️ 现场问题描述 (大白话)：", height=220, placeholder="描述现场看到的真实情况...")
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown('<div class="main-btn">', unsafe_allow_html=True)
                
                if st.button("🚀 使用 AI 润色并打分"):
                    if not raw_finding:
                        st.error("请先输入现场问题描述。")
                    else:
                        df = st.session_state['vda_data']
                        q_text, b_points, a_hints = "未知", "未知", "未知"
                        if "reference" in df.columns:
                            selected_row = df[df['reference'].astype(str) == str(clause)]
                            if not selected_row.empty:
                                if "Question" in df.columns: q_text = selected_row['Question'].values[0]
                                if "Bullet Points" in df.columns: b_points = selected_row['Bullet Points'].values[0]
                                if "Additional Hints and Requirements for PT ISPA" in df.columns: a_hints = selected_row['Additional Hints and Requirements for PT ISPA'].values[0]
                        
                        prompt = f"""
                        你是一位资深 VDA 6.3 主任审核员。条款:{clause}。要求:{q_text}。参考:{b_points}。附加:{a_hints}。
                        现场白话描述：“{raw_finding}”
                        请严格按两部分输出：
                        第一部分（直接输出润色后的Finding）：以“在审核期间观察到...”或“现场发现...”开头，用专业术语描述，150字内。
                        第二部分（以加粗的“系统降级与扣分建议：”开头）：给出建议分数（如：建议 4 分）和简单理由。
                        """
                        with st.spinner("🤖 DeepSeek 引擎正在分析中..."):
                            try:
                                response = client.chat.completions.create(
                                    model="deepseek-chat",
                                    messages=[
                                        {"role": "system", "content": "你是一个专业的汽车行业质量审核助手。"},
                                        {"role": "user", "content": prompt}
                                    ],
                                    temperature=0.3 
                                )
                                full_answer = response.choices[0].message.content
                                if "系统降级与扣分建议：" in full_answer:
                                    parts = full_answer.split("系统降级与扣分建议：")
                                    st.session_state['ai_finding_result'] = parts[0].strip()
                                    st.session_state['ai_score_result'] = "系统降级与扣分建议：\n" + parts[1].strip()
                                else:
                                    st.session_state['ai_finding_result'] = full_answer
                                    st.session_state['ai_score_result'] = "无法提取独立打分，请参考上方描述。"
                                st.session_state['ai_triggered'] = True
                            except Exception as e:
                                st.error(f"❌ 调用 DeepSeek 失败：{e}")
                st.markdown('</div>', unsafe_allow_html=True)

            with col2:
                st.markdown("#### 🤖 第二步：AI 分析结果")
                if st.session_state.get('ai_triggered', False):
                     st.info(f"**✅ AI 专业润色后的 Finding：**\n\n{st.session_state['ai_finding_result']}")
                     st.warning(f"**⚠️ AI 智能评分建议：**\n\n{st.session_state['ai_score_result']}")
                     st.markdown('<div class="main-btn">', unsafe_allow_html=True)
                     if st.button("💾 确认此问题并前往 5-Why 分析"):
                         st.session_state['current_page'] = "page3"
                         st.rerun()
                     st.markdown('</div>', unsafe_allow_html=True)

elif menu == "page3":
    st.title("📑 AI 辅助 5-WHY 问题原因分析")
    st.markdown("通过追问为什么，探究系统性失效的根本原因。")
    
    finding = st.session_state.get('ai_finding_result', "")
    if not finding:
        st.warning("⚠️ 系统中尚未确认任何审核发现。请先在 [录入与 AI 分析] 页面确认一个问题。")
    else:
        with st.container():
            st.markdown("##### 📍 表面现象 (The Phenomenon)")
            why_1 = st.text_area("Why 1: 发生了什么问题？(已自动带入)", value=finding, height=100)
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown('<div class="main-btn">', unsafe_allow_html=True)
            
            if st.button("🧠 让 AI 专家自动推导后续 4 个 Why"):
                prompt_5why = f"""
                你是一位质量工程师，精通 5-WHY 分析。现象（第1个Why）是：“{why_1}”
                请推导出后续的 4 个 Why，指向系统失效。严格按格式输出：
                Why 2: [推导]
                Why 3: [推导]
                Why 4: [推导]
                Why 5 (根本原因): [推导]
                """
                with st.spinner("🤖 DeepSeek 正在进行逻辑推理..."):
                    try:
                        response = client.chat.completions.create(
                            model="deepseek-chat",
                            messages=[{"role": "user", "content": prompt_5why}],
                            temperature=0.5 
                        )
                        st.session_state['five_why_result'] = response.choices[0].message.content
                    except Exception as e:
                        st.error(f"❌ 调用 DeepSeek 失败：{e}")
            st.markdown('</div>', unsafe_allow_html=True)

        if st.session_state.get('five_why_result'):
            st.markdown("---")
            st.markdown("##### 📍 逻辑推导结果")
            st.success(st.session_state['five_why_result'])
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("##### 📥 报告导出")
            
            col1, col2 = st.columns(2)
            with col1: person_input = st.text_input("责任人", placeholder="输入姓名...")
            with col2: deadline_input = st.text_input("完成期限", placeholder="YYYY-MM-DD")
            
            # 生成 Word 文件流
            word_file = create_word_report(
                st.session_state['ai_finding_result'],
                st.session_state['five_why_result'],
                person_input,
                deadline_input
            )
            
            st.markdown("<br>", unsafe_allow_html=True)
            st.download_button(
                label="📥 下载 5-Why 原因分析报告 (Word格式)",
                data=word_file,
                file_name="VDA6.3_5Why_分析报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
